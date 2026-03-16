from __future__ import annotations
from pathlib import Path
from voice_rag.core.models import Document

class DocxLoader:
    supported_extensions = [".docx"]
    def load(self, path: Path) -> Document:
        try:
            from docx import Document as DocxDocument
        except ImportError as e:
            raise ImportError("python-docx is required for DOCX parsing. Install with: pip install voice-rag[docx]") from e
        doc = DocxDocument(str(path))
        return Document(source=str(path), content="\n".join(p.text for p in doc.paragraphs))
